import kivy
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.recycleview import RecycleView
from kivy.uix.recycleview.views import RecycleDataViewBehavior
from kivy.uix.label import Label
from kivy.uix.recycleboxlayout import RecycleBoxLayout
from kivy.uix.behaviors import FocusBehavior
from kivy.uix.recycleview.layout import LayoutSelectionBehavior
from kivy.properties import ObjectProperty, BooleanProperty, ListProperty
from kivy.uix.popup import Popup
from kivy.uix.button import Button

import sqlite3
import docx
import os, subprocess


class Librarian(object):

    def __init__(self):
        self.books = None
        self.update()

    def update(self):
        """Updates the data."""
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        c.execute("SELECT * FROM library")
        data = c.fetchall()
        conn.close()

        self.books = data

    def store_book(self, title, author, genre):
        """Stores the book in the database"""
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        c.execute("INSERT INTO library VALUES (?, ?, ?)", (title, author, genre))
        conn.commit()
        conn.close()

        self.update()

    def delete_book(self, titulo):
        """Deletes from the database."""
        try:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("DELETE FROM library WHERE title = '%s'" %titulo)
            conn.commit()
            conn.close()

            self.update()
        except:
            print("Didn't work")

class DeletePopup(Popup):
    pass

class TextInputPopup(Popup):
    title_text_input = ObjectProperty()
    author_text_input = ObjectProperty()
    genre_text_input = ObjectProperty()
    libr = ObjectProperty(Librarian())

    def adding_book(self):
        """Adds the book to the database and the display data in the RecycleView."""
        self.libr.store_book(self.title_text_input.text, self.author_text_input.text, self.genre_text_input.text) # adds to database
        App.get_running_app().root.ids.rv.update_data(main_instance.order_by_author) # adds to recycleview data

class SelectableRecycleBoxLayout(FocusBehavior, LayoutSelectionBehavior,
                                RecycleBoxLayout):
    '''Adds selection and focus behaviour to the view.'''

class SelectableLabel(RecycleDataViewBehavior, Label):
    '''Add selection support to the Label'''
    index = None
    selected = BooleanProperty(False)
    selectable = BooleanProperty(True)
    libr = ObjectProperty(Librarian())
    data = None

    def refresh_view_attrs(self, rv, index, data):
        """Catch and handle the view changes."""
        self.index = index
        self.data = data
        return super(SelectableLabel, self).refresh_view_attrs(rv, index, data)

    def on_touch_down(self, touch):
        """Add selection on touch down."""
        if super(SelectableLabel, self).on_touch_down(touch):
            return True
        if self.collide_point(*touch.pos) and self.selectable:
            return self.parent.select_with_touch(self.index, touch)

    def apply_selection(self, rv, index, is_selected):
        """Respond to the selection of items in the view."""
        self.selected = is_selected

        # Checks if label has been selected to put in selected_data list property in 'rv'.
        # If it isn't selected or it has been deselected it'll check if it's in the selected_data list
        # so it can remove it if it is.
        if self.selected:
            rv.selected_data.append(rv.data[index])
        else:
            # If data is in selected list
            if self.data in rv.selected_data:
                rv.selected_data.remove(self.data)


class LibraryRV(RecycleView):
    libr = ObjectProperty(Librarian())
    selected_data = ListProperty(list())

    def __init__(self, **kwargs):
        super(LibraryRV, self).__init__(**kwargs)
        self.update_data()

    def update_data(self, author=False):
        """Decides how the data will sorted to be displayed."""
        self.libr.update() # always updates gets the latest data from database.
        if author: # will sort data alphabetically by author
            self.libr.books.sort(key=lambda t: t[1])
            self.data = [{
                'data_index': self.libr.books.index(book),
                'title': book[0],
                'genre': book[2],
                'text': book[0]+': '+book[1],
                }
                for book in self.libr.books
            ]
        else:
            # By default, sort data alphabetically by title
            self.libr.books.sort(key=lambda t: t[0])
            self.data = [{
                'data_index': self.libr.books.index(book),
                'title': book[0],
                'genre': book[2],
                'text': book[0]+': '+book[1],
                }
                for book in self.libr.books
            ]

    def delete_selections(self):
        """Deletes all books that have been selected."""
        if self.selected_data:
            for data_value in self.selected_data:
                title = data_value['title']

                self.data.remove(data_value)

                self.libr.delete_book(title)

            # After deletion of all selected items,
            # reset all values; the selected layout, selected data, reconstruct new data
            self.layout_manager.clear_selection()
            self.selected_data = list()
            self.data = [{
                'data_index': i,
                'title': d['title'],
                'genre': d['genre'],
                'text': d['text']
                }
                for i, d in enumerate(self.data)
            ]
        else:
            # Creates and error pop if no books have been selected.
            popup = DeletePopup()
            popup.open()


class ShelfLayout(BoxLayout):
    order_by_author = BooleanProperty(False)
    libr = ObjectProperty(Librarian())

    def __init__(self, **kwargs):
        super(ShelfLayout, self).__init__(**kwargs)

    def add_book(self):
        """Creates a PopUp instance that will handle adding the book."""
        popup = TextInputPopup()
        popup.open()

    def print_books(self):
        """Creates a word document with all the books in the database."""
        # Creates a dict() for the Word doc based on the rv.data
        # dict() skeleton : {'genre': ['Title: Author', 'Title: Author']}
        display_data = dict()
        for i in self.ids.rv.data:
            key_created = display_data.get(i['genre'])
            print(key_created)
            if key_created:
                key_created.append(i['text'])
            else:
                display_data[i['genre']] = [i['text']]

        # Will create a Word doc and display the data
        # alphabetically by genre.
        doc = docx.Document()
        for k in sorted(display_data.keys()):
            doc.add_heading(k, 3)
            for v in display_data[k]:
                doc.add_paragraph(v)

        doc.save('library.docx')
        subprocess.call(('open', 'library.docx')) # opens word doc


class MyApp(App):

    def build(self):
         return ShelfLayout()

if __name__ == "__main__":
    MyApp().run()
